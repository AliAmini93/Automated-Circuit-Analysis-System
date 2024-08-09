import logging

logger = logging.getLogger('logger_Info')

import docx
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from ic_info.path import Path_Datasheets, Path_Report, Path_Input


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    r = paragraph.add_run()
    r._r.append(hyperlink)

    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def create_doc(data, pictures_IC):
    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    style1 = doc.styles.add_style('rtl', WD_STYLE_TYPE.PARAGRAPH)
    style1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    style2 = doc.styles.add_style('center', WD_STYLE_TYPE.PARAGRAPH)
    style2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    data = data

    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row = table.rows[0].cells

    row[0].text = 'لینک دیتاشیت'
    row[1].text = 'مشخصات فنی'
    row[2].text = 'سازنده'
    row[3].text = 'دسته'
    row[4].text = 'تصویر'
    row[5].text = 'شناسه تراشه'
    row[6].text = 'شماره قطعه'
    row[7].text = 'شماره'

    id = 1
    for ic, pic in zip(data, pictures_IC):

        row = table.add_row().cells
        row[1].text = ic['Description']
        row[2].text = ic['Manufacturer']
        row[3].text = ic['Category']

        paragraph = row[4].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(Path_Input + pic, width=900000, height=900000)  # ic['picture']

        row[5].text = ic['textOCR']
        row[6].text = ic['ManufacturerPartNumber']
        row[7].text = str(id)
        id += 1

        paragraph = row[0].paragraphs[0]

        if '_id' not in ic.keys() or ic['DataSheetUrl'] == '-':
            add_hyperlink(paragraph, '', '')
        else:
            if ic['DataSheetUrl'] == '':
                add_hyperlink(paragraph, str(ic['_id']) + '.pdf',
                              Path_Datasheets + str(ic['_id']) + '.pdf')
            else:
                add_hyperlink(paragraph, ic['DataSheetUrl'], ic['DataSheetUrl'])  # ic['DataSheetUrl']

        run = paragraph.runs

        font = run[0].font
        font.color.rgb = RGBColor(0, 0, 255)
        paragraph = row[0].add_paragraph()

    table.style = 'Colorful List'

    for cell in table.rows[0].cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            paragraph.style = 'center'
            for run in paragraph.runs:
                font = run.font
                font.name = 'B Nazanin'
                font.complex_script = True
                font.size = Pt(12)

    for row in table.rows[1:]:
        for i, cell in enumerate(row.cells):
            paragraphs = cell.paragraphs
            if i in [6, 5, 3, 2]:
                for paragraph in paragraphs:
                    paragraph.style = 'center'
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(10)
            if i == 6:
                for paragraph in paragraphs:
                    if paragraph.text == 'پیدا نشد':
                        for run in paragraph.runs:
                            font = run.font
                            font.complex_script = True
            if i == 7:
                for paragraph in paragraphs:
                    paragraph.style = 'center'
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'B Nazanin'
            if i in [1, 0]:
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(10)

    table.autofit = False
    table.allow_autofit = False
    widths = (Inches(1.5), Inches(1.5), Inches(0.8), Inches(0.8), Inches(1), Inches(1), Inches(1), Inches(0.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.save(Path_Report + 'Report.docx')
